"""
Preprocessing layer: extract clinical text from uploaded files (PDF, DOCX, images).
Output is plain text only; no pipeline or agent logic. Used by the dashboard before run_pipeline().
"""

import io
import os
import re
import subprocess
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

# Project root (parent of src/)
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".png", ".jpg", ".jpeg")
SUPPORTED_MIME_PDF = "application/pdf"
SUPPORTED_MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_MIME_IMAGES = ("image/png", "image/jpeg", "image/jpg")


def _extract_pdf_pypdf(file_bytes: bytes) -> Optional[str]:
    """Extract text using pypdf. Returns None on failure or empty."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            try:
                t = page.extract_text()
                if t:
                    parts.append(t)
            except Exception:
                continue
        text = "\n".join(parts).strip()
        return text if text else None
    except Exception:
        return None


def _extract_pdf_pymupdf(file_bytes: bytes) -> Optional[str]:
    """Extract text using PyMuPDF (fitz). Works on many PDFs where pypdf fails."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            t = page.get_text()
            if t and t.strip():
                parts.append(t.strip())
        doc.close()
        text = "\n".join(parts).strip()
        return text if text else None
    except Exception:
        return None


def extract_text_from_pdf(file_bytes: bytes) -> Optional[str]:
    """Extract text from PDF. Tries pypdf first, then PyMuPDF. Returns None on failure."""
    first = _extract_pdf_pypdf(file_bytes)
    if first and len(first.strip()) >= 20:
        return first.strip()
    fallback = _extract_pdf_pymupdf(file_bytes)
    if fallback and len(fallback.strip()) >= 20:
        return fallback.strip()
    return (first.strip() if first and first.strip() else None)


def _extract_docx_python_docx(file_bytes: bytes) -> Optional[str]:
    """Extract text from .docx using python-docx (paragraphs + tables)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts).strip()
        return text if text else None
    except Exception:
        return None


def _extract_docx_docx2txt(file_bytes: bytes) -> Optional[str]:
    """Extract text using docx2txt (often works when python-docx misses content)."""
    try:
        import docx2txt
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            text = docx2txt.process(tmp.name)
        return text.strip() if text and text.strip() else None
    except Exception:
        return None


def _extract_docx_raw_zip(file_bytes: bytes) -> Optional[str]:
    """Extract text from .docx by reading word/document.xml (ZIP + XML). No extra deps; works when other libs fail."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as z:
            if "word/document.xml" not in z.namelist():
                return None
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
        # OOXML namespaces
        ns = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        }
        parts = []
        for t in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
            if t.text and t.text.strip():
                parts.append(t.text.strip())
            if t.tail and t.tail.strip():
                parts.append(t.tail.strip())
        text = " ".join(parts)
        text = re.sub(r"\s+", " ", text).strip()
        return text if text else None
    except Exception:
        return None


def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    """Extract text from Word document. Tries python-docx, docx2txt, then raw ZIP/XML."""
    first = _extract_docx_python_docx(file_bytes)
    if first and len(first.strip()) >= 10:
        return first.strip()
    fallback = _extract_docx_docx2txt(file_bytes)
    if fallback and len(fallback.strip()) >= 10:
        return fallback.strip()
    raw = _extract_docx_raw_zip(file_bytes)
    if raw and len(raw.strip()) >= 10:
        return raw.strip()
    for candidate in (first, fallback, raw):
        if candidate and candidate.strip():
            return candidate.strip()
    return None


def _run_tesseract_subprocess(tesseract_cmd: str, image_path: str) -> tuple[Optional[str], Optional[str]]:
    """Run tesseract.exe via subprocess; returns (text, error_message)."""
    try:
        # Use "-" so tesseract writes text to stdout
        result = subprocess.run(
            [tesseract_cmd, image_path, "-", "quiet"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
            creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0,
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout.strip(), None
        if result.stderr:
            return None, result.stderr.strip() or "Tesseract failed."
        return None, "Tesseract returned no text."
    except subprocess.TimeoutExpired:
        return None, "Tesseract timed out."
    except FileNotFoundError:
        return None, f"Tesseract not found at: {tesseract_cmd}"
    except Exception as e:
        return None, str(e)


def extract_text_from_image(file_bytes: bytes) -> tuple[Optional[str], Optional[str]]:
    """
    Extract text from image using OCR. Returns (text, custom_error_message).
    Set TESSERACT_CMD in .env to the full path to tesseract.exe if it is not on system PATH.
    """
    # Load .env so TESSERACT_CMD is available (dashboard may not have loaded it yet)
    try:
        from dotenv import load_dotenv
        load_dotenv(_PROJECT_ROOT / ".env")
    except Exception:
        pass
    tesseract_cmd = os.environ.get("TESSERACT_CMD", "").strip()
    if tesseract_cmd:
        tesseract_cmd = os.path.normpath(os.path.expanduser(tesseract_cmd))
        # If path is a directory, use tesseract.exe inside it
        if os.path.isdir(tesseract_cmd):
            tesseract_cmd = os.path.join(tesseract_cmd, "tesseract.exe")
    if tesseract_cmd and os.path.isfile(tesseract_cmd):
        # Prefer subprocess so we don't rely on pytesseract finding the executable
        suffix = ".png"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp.flush()
            tmp_path = tmp.name
        try:
            text, err = _run_tesseract_subprocess(tesseract_cmd, tmp_path)
            if text:
                return (text if text.strip() else None, None)
            if err:
                return None, (
                    "Image OCR failed. Install Tesseract OCR for image support: "
                    "https://github.com/tesseract-ocr/tesseract. "
                    f"Detail: {err}"
                )
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
    elif tesseract_cmd:
        return None, (
            f"Tesseract not found at path in TESSERACT_CMD. Check .env: {tesseract_cmd}"
        )
    # Fallback: pytesseract (when TESSERACT_CMD not set, e.g. Tesseract on PATH)
    try:
        import pytesseract
        from PIL import Image
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode not in ("L", "RGB", "RGBA"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img).strip()
        return (text if text else None, None)
    except Exception as e:
        err = str(e).lower()
        if "tesseract" in err or "not found" in err or "no such file" in err or "ocr" in err:
            return None, (
                "Image OCR failed. Install Tesseract OCR for image support: "
                "https://github.com/tesseract-ocr/tesseract"
            )
        return None, None


def get_file_extension(filename: str) -> str:
    """Return lowercase extension including dot, e.g. .pdf."""
    if not filename:
        return ""
    p = filename.rsplit(".", 1)
    return ("." + p[-1].lower()) if len(p) == 2 else ""


def is_supported_file(filename: str) -> bool:
    """True if filename has a supported extension."""
    ext = get_file_extension(filename)
    return ext in SUPPORTED_EXTENSIONS


def extract_text_from_file(file_bytes: bytes, filename: str) -> tuple[Optional[str], Optional[str]]:
    """
    Extract text from uploaded file based on extension.
    Returns (extracted_text, error_message).
    If success: (text, None). If failure: (None, "Unable to extract...") or (None, "Unsupported file type...").
    """
    ext = get_file_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        return None, "Unsupported file type. Please upload PDF, DOCX, PNG, or JPG."

    text = None
    custom_error = None
    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        text = extract_text_from_docx(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg"):
        text, custom_error = extract_text_from_image(file_bytes)

    if text is None or not text.strip():
        return None, (custom_error or "Unable to extract clinical text from this file.")
    return text.strip(), None
